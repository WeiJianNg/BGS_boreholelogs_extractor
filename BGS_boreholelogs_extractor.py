# Author: Ng Wei Jian 
# Title: Automating Extraction of BGS borehole logs 
# The purpose of this script is to allow extraction of borehole logs from BGS
import re 
import urllib.request
import xlsxwriter
import os
import docx
import docx.shared
from docx.enum.section import WD_ORIENT
# ==============
# Display Header
# ==============
print(80*'=')
print('Author: Ng Wei Jian')
print('Title: BGS Borehole Logs Extractor')
print(80*'=')
# ==============
#      End 
# ==============

# =============================================================
# Directory Checker - creates folder/director if does not exist
# =============================================================
def check_dir(file_name):
    directory = os.path.dirname(file_name)
    if not os.path.exists(directory):
        os.makedirs(directory)
# ==============
#      End 
# ==============
# =============================================================
# Check input folder is provided
# =============================================================
def check_input(file_name):
    if not os.path.isfile(file_name):
        error_message = input('ERROR! - GeoIndexData.txt - Input file does not exist in folder, please check folder / name of input file')
        exit()
# ==============
#      End 
# ==============
check_input("GeoIndexData.txt")
# Create folder to store images
check_dir(os.path.join('Borehole_logs',''))

# =============================================================
# Reading, Cleaning and storing data exported from BGS website
# =============================================================
BGS = open("GeoIndexData.txt")
# Read first line in data which is the header
headings = BGS.readline().split()
# Create Excel workbook to store data
excel = xlsxwriter.Workbook('Data.xlsx')
worksheet = excel.add_worksheet('My_data')
column_counter = 0
for k in headings:
    worksheet.write(0, column_counter, k)
    column_counter += 1
# Headings contain the headings of the data
# [1]: Record; [2]: Reference; [3]: Name; [4]: Borehole Length; [5]: Year known; [6]: Eastings; [7]: Northings
data = []
for line in BGS:
    data.append(line.split("\t"))

# Cleaning up the link column using a function

def clean(lt):
    string = lt[0]
    m = re.match("<a href='(.*)\' class", string)
    new_lt = [m.group(1)+'/images/'] + lt[1::]
    return new_lt

data = list(map(clean, data))
data1 = []
# filtering out the confidential data -> This is inefficient
# Writing to error log for files that needs to be purchased
print('Reading Data \nDetecting borehole logs that needs to be purchased.')
count_purchase = 0
error_log = open("error_log.txt","w+")
for i in data:
    if "shop" in i[0]:
        print('***Note: Borehole ' + i[1] + ' needs to be purchased. See error log.')
        error_log.write('***Note: Borehole ref: ' + i[1] + ' needs to be purchased from the website. Link: http://shop.bgs.ac.uk/GeoRecords \n')
        count_purchase += 1
    else:
        data1.append(i)
    
if count_purchase == 0:
    print('None found.\n')
else:
    print('Total of ' + str(count_purchase) + ' logs needs to be purchased. see error_log.\n')

# In cases where the name of the BH references has special character; this step will remove it
def remove_special(lt):
    lt_new = lt
    lt_new[1] = re.sub('/','',lt[1])
    return lt_new
# cleaned data - ready to be used for extracting the images from the website
data1 = list(map(remove_special,data1))
# ==============
#      End 
# ==============

# ================================================================================
# Writing functions to extract images from the website using Python urllib library
# ================================================================================
# Assigning agent for opener to avoid 403 error 
opener = urllib.request.build_opener()
opener.addheaders = [('User-Agent', 'Mozilla/5.0')]
urllib.request.install_opener(opener)

# Function to clean the url list contained in data1[0] / images folder
def clean2 (my_string):
    my_string = list(my_string)
    my_string.pop(0) # pop the b
    my_string.pop(0)
    for i in range(5):
        my_string.pop(-1)
    return "".join(my_string)+'.png'

# Function to download / retrive images
def download_image(url, file_path, file_name):
    full_path = file_path + file_name + '.png'
    print('Downloading Borehole ' + i[1] + ' image no. ' + str(count)+' ......', end =" "),
    urllib.request.urlretrieve(url, full_path)

# ==============
#      End 
# ==============

# ===============================
# Defining format for output word
# ===============================
document = docx.Document()
paragraph_format = document.styles['Normal'].paragraph_format
paragraph_format.space_before = docx.shared.Pt(0)
paragraph_format.space_after = docx.shared.Pt(0)
section = document.sections[-1]

# Function to rescale image
def scale_image_landscape(picture, max_height):
    current_height = docx.shared.Inches(picture.height)
    current_width = docx.shared.Inches(picture.width)
    # scale to fit image to max_width
    scale_height = max_height / current_height
    picture.height = int(current_height*scale_height)
    picture.width = int(current_width*scale_height)


def scale_image_portrait(picture, max_width, max_height):
    current_height = docx.shared.Inches(picture.height)
    current_width = docx.shared.Inches(picture.width)
    # scale to fit image to max_width
    scale_width = max_width / current_width
    if current_height*scale_width > max_height:
        scale_height = max_height / current_height
        picture.height = int(current_height*scale_height)
        picture.width = int(current_width*scale_width)
    else:
        picture.height = int(current_height*scale_width)
        picture.width = int(current_width*scale_width)
        
# ==============
#      End 
# ==============

counter = 0 # this is for tracking progress
new_length = len(data1)
row_counter = 1
image_count = 0
column_counter = 0
count1 = 0 # this is to check amount of file not downloaded
count_error = 0 # this is to check amount of file not downloaded\

for i in data1:
    print('*****Initializing*****Borehole '+i[1] +' download')
    print('Retrieving download destination for Borehole '+i[1])
    url = list(urllib.request.urlopen(i[0]))
    url = list(map(str, url))
    url = list(map(clean2, url))
    count = 1 # this is to increment the file
    for j in url:
        count1 += 1
        try:
            # Downloading image
            download_image(j, 'Borehole_logs/', i[1]+'_'+ str(count))
            print('Completed')
            image_count += 1
            # Writing to word document
            print('Writing to Word Document ......', end=" ")
            image_loc = os.path.join('Borehole_logs/', i[1]+'_'+ str(count)+'.png')
            document.add_paragraph('Reference : ' + i[1])
            document.add_paragraph('Project Name : ' + i[2])
            document.add_paragraph('Year : ' + i[4])                      
            document.add_paragraph('Eastings, Northings : '+ i[5] + ', ' + i[6].rstrip())
            document.add_paragraph('Sheet ' + str(count) + ' of '+str(len(url)))
            picture = document.add_picture(image_loc)
            if picture.height > picture.width:
                if section.orientation != WD_ORIENT.PORTRAIT:
                    new_width, new_height = section.page_height, section.page_width          
                    section.orientation = WD_ORIENT.PORTRAIT
                    section.page_width = new_width
                    section.page_height = new_height
                scale_image_portrait(picture, docx.shared.Inches(5.5), docx.shared.Inches(7.2))
                document.add_paragraph('')
                document.add_section()
                count += 1
            else:
                if section.orientation != WD_ORIENT.LANDSCAPE:
                    new_width, new_height = section.page_height, section.page_width               
                    section.orientation = WD_ORIENT.LANDSCAPE
                    section.page_width = new_width
                    section.page_height = new_height
                scale_image_landscape(picture, docx.shared.Inches(4.8))
                document.add_paragraph('')
                document.add_section()
                count += 1
            document.save('Borehole_logs.docx')
            print('completed')

            
        except urllib.error.HTTPError as e:
            if e.code in (..., 403, ...):
                print('Error downloading Borehole ' + i[1] + ' image no. ' + str(count) + ' due to server error. Check Error log')
                error_log.write('403 Error downloading Borehole ' + i[1] + ' image no. ' + str(count) + '. link: ' + j + '\n')
                count_error += 1
                # appending the error link to error list for checking
                continue
            
    for k in i:
        worksheet.write(row_counter, column_counter, k)
        column_counter += 1
        

    column_counter = 0
    row_counter += 1
    count = 0
    counter += 1
    print('*********End**********Borehole '+i[1] +' Download Completed')
    print('Overall Progress = '+ str(round(100*counter/new_length,1)) +'%\n')

# ==== End ====
# closing all files and saving word document

error_log.close()  # closing error log file
excel.close() # closing excel file

# ===============
# Display summary
# ===============
print(80*'=')
print('Summary Report')
print(80*'=')
print('Total number of logs in area:', len(data))
print('Number of logs that needs to be purchased: %d ' % (int((len(data)-len(data1)))))
print('Number of images downloaded: %d ' % (image_count))
print('Number of images not downloaded due to 403 error: ', count_error, '. Please check error log for more info')
# ===============
#       End
# ===============
k = input("press close to exit")

    

